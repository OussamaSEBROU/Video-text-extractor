import streamlit as st
import os
import tempfile
from pathlib import Path
import io
from docx import Document
import google.generativeai as genai
import time # For time.sleep during AI file processing

# --- Configuration ---
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    st.error("AI service key not found. Please set the 'GEMINI_API_KEY' environment variable.")
    st.stop()

genai.configure(api_key=GEMINI_API_KEY)

st.set_page_config(
    page_title="TahiriExtractor - Video Ultra Transcription",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- Helper Function for AI Video Processing ---

def extract_text_with_ai(video_file_path):
    """
    Extracts comprehensive visual content and on-screen text from a video
    using an advanced AI model. This function aims for a "transcription-like"
    output based purely on visual observation.

    Important Note on Transcription:
    This model excels at understanding visual content and reading on-screen text.
    It **does NOT perform direct audio-to-text transcription**. The output will be
    a detailed descriptive summary of the video's visual elements and any text
    appearing in frames, structured to resemble a continuous narrative or report.
    For true spoken word transcription, a dedicated ASR service is required.
    """
    st.info("Initiating video analysis with our AI for visual content extraction. "
            "Please note: This process focuses on what is *visually observable* in the video, "
            "including any on-screen text, and infers narrative from visual cues. "
            "It does NOT perform audio-to-text transcription.")

    uploaded_file_name = None
    try:
        with st.spinner("Uploading video for AI analysis..."):
            video_file = genai.upload_file(video_file_path)
            uploaded_file_name = video_file.name

        processing_bar = st.progress(0, text="Processing video with AI... Please wait, this may take a moment.")
        progress_percentage = 0
        while video_file.state.name == "PROCESSING":
            time.sleep(2)
            video_file = genai.get_file(uploaded_file_name)
            progress_percentage = min(progress_percentage + 5, 99)
            processing_bar.progress(progress_percentage)

        if video_file.state.name == "FAILED":
            st.error("Failed to process video with our AI. Please try again or with a different video.")
            return "Error: Video processing failed."

        # --- Refined Prompt for Direct Content Extraction (UNMODIFIED AS PER INSTRUCTION) ---
        prompt = (
            "You are a highly precise content extraction AI. Your task is to extract the text from the video without change anything thing in the content"
" the language of extracted text should be same of the video .. be carefuly on that"

            "Present the extracted content in clean, well-formatted paragraphs, ordering events chronologically as they appear in the video. "
           
        )
        
        # Using gemini-2.0-flash model (UNMODIFIED AS PER INSTRUCTION)
        model = genai.GenerativeModel('gemini-2.0-flash')
        response = model.generate_content([prompt, video_file], stream=False)
        
        if response and hasattr(response, 'text'):
            return response.text
        else:
            st.error("Our AI service did not return expected text content.")
            return "Error: AI analysis did not return valid text."

    except Exception as e:
        st.error(f"An error occurred during AI analysis: {e}")
        return "Error: Could not extract information using our AI service."
    finally:
        if uploaded_file_name:
            try:
                genai.delete_file(uploaded_file_name)
                st.success("Temporary file successfully cleaned up from our AI system.")
            except Exception as e:
                st.warning(f"Failed to delete uploaded file from our AI system: {e}. "
                           "Please check your usage or manually clear if possible.")

# --- Streamlit UI Layout ---

with st.sidebar:
    st.header("How to Use TahiriExtractor")
    st.markdown("""
    * **Step 1:** Upload your video file using the "Upload a video file" button on the main page.
    * **Step 2:** Click the "Generate Content Summary" button.
    * **Step 3:** Our AI-powered system will analyze the visual content of your video.
    * **Step 4:** The extracted descriptive text (visual 'transcription') will be displayed in the main area, with options to copy or download as a Word file.
    """)

    st.header("About TahiriExtractor")
    st.markdown("""
    **TahiriExtractor** is an innovative application leveraging **advanced Artificial Intelligence** to:
    * Extract deep textual insights from video visual content.
    * Analyze video frames for a comprehensive, structured summary.
    * Capture and transcribe any on-screen text, actions, and visual narratives.
    * **Important Note:** This application does not transcribe audio content (speech-to-text).
    """)

    st.header("Contact Us")
    st.markdown("""
    Have questions, feedback, or need support? We'd love to hear from you!

    Contact us via email:
    [TahiriExtractor.veo.net](mailto:oussama.sebrou@gmail.com?subject=Inquiry%20from%20TahiriExtractor%20App&body=Hello%20TahiriExtractor%20Team%2C%0A%0AI%20am%20contacting%20you%20regarding%20...)
    """)

# Main title with blue color
st.markdown("<h1 style='color: #1E90FF; text-align: center;'>🎥 TahiriExtractor - Video Ultra Transcription</h1>", unsafe_allow_html=True) # Updated title with blue color and centered

st.markdown("""
<div style='background-color: #f0f2f6; padding: 20px; border-radius: 10px; text-align: center;'>
    <p style='font-size: 1.1em; color: #333;'>
        Welcome to <strong>TahiriExtractor</strong>, your ultimate tool for extracting comprehensive visual information from videos.
        This application harnesses the power of <strong>advanced Artificial Intelligence</strong> to analyze your video content,
        identifying and transcribing on-screen text, objects, and actions to provide a detailed, readable summary.
        Perfect for researchers, content creators, and anyone needing to quickly grasp the visual narrative of a video.
    </p>
    <p style='font-size: 1em; color: #555;'>
        <strong>Key Point:</strong> We support videos up to <strong>15 minutes</strong> in duration for efficient analysis.
        Our system focuses on visual analysis and does not perform audio-to-text transcription.
    </p>
</div>
<br>
""", unsafe_allow_html=True) # Professional, concise introduction

uploaded_file = st.file_uploader(
    "Upload a video file (MP4, MOV, MKV, AVI, WEBM, etc.)",
    type=["mp4", "mov", "mkv", "avi", "webm"],
    help="Our AI system supports videos up to approximately 15 minutes in length. Larger files may fail or take longer."
)

transcript_display_area = st.empty()

if uploaded_file is not None:
    # Save the uploaded file to a temporary location
    with tempfile.NamedTemporaryFile(delete=False, suffix=Path(uploaded_file.name).suffix) as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        video_path = tmp_file.name

    st.video(video_path)

    if st.button("Generate Content Summary", type="primary", use_container_width=True):
        with st.spinner("Analyzing video visual content with our AI... This may take a moment based on video length and complexity."):
            extracted_text = extract_text_with_ai(video_path)

        with transcript_display_area.container():
            st.markdown("---")
            st.subheader("Extracted Visual Content (Visual 'Transcription')")
            
            # Display the extracted text in a non-editable, formatted way (like a chatbot response)
            st.markdown(
                f"""
                <div style='background-color: #e6f7ff; padding: 15px; border-radius: 10px; border-left: 5px solid #1E90FF; margin-bottom: 15px; overflow-wrap: break-word;'>
                    {extracted_text}
                </div>
                """,
                unsafe_allow_html=True
            )

            # Copyable text area (with built-in copy icon)
            st.text_area(
                "Copyable Visual Content Summary",
                value=extracted_text,
                height=300,
                key="copy_summary_area",
                help="You can easily copy the entire summary from this text box using the built-in copy button."
            )

            # Create and download Word document
            doc = Document()
            doc.add_heading('Video Visual Content Summary', level=1)
            for paragraph_text in extracted_text.split('\n\n'):
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())

            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)

            st.download_button(
                label="Download as Word (DOCX)",
                data=bio.getvalue(),
                file_name="video_visual_content_summary.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Download the extracted visual content summary as a Microsoft Word document for offline use."
            )
    
    # Ensure temporary file is cleaned up after use
    if os.path.exists(video_path):
        os.unlink(video_path)

