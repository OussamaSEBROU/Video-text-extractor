import streamlit as st
import os
import tempfile
from pathlib import Path
import io
from docx import Document
import google.generativeai as genai
import time # For time.sleep during AI file processing

# --- Configuration ---
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    st.error("AI service key not found. Please set the 'GEMINI_API_KEY' environment variable.")
    st.stop()

genai.configure(api_key=GEMINI_API_KEY)

st.set_page_config(
    page_title="TahiriExtractor - Video Ultra Transcription",
    layout="wide",
    initial_sidebar_state="expanded",
)

# --- Initialize Session State for Chat ---
if 'extracted_text' not in st.session_state:
    st.session_state.extracted_text = ""
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'main_page_selection' not in st.session_state: # Renamed to avoid confusion with sidebar elements
    st.session_state.main_page_selection = "Video Extraction" # Default main page

# --- Helper Function for AI Video Processing ---

def extract_text_with_ai(video_file_path):
    """
    Extracts comprehensive visual content and on-screen text from a video
    using an advanced AI model. This function aims for a "transcription-like"
    output based purely on visual observation.

    Important Note on Transcription:
    This model excels at understanding visual content and reading on-screen text.
    It **does NOT perform direct audio-to-text transcription**. The output will be
    a detailed descriptive summary of the video's visual elements and any text
    appearing in frames, structured to resemble a continuous narrative or report.
    For true spoken word transcription, a dedicated ASR service is required.
    """
    st.info("Initiating video analysis with our AI for visual content extraction. "
            "Please note: This process focuses on what is *visually observable* in the video, "
            "including any on-screen text, and infers narrative from visual cues. "
            "It does NOT perform audio-to-text transcription.")

    uploaded_file_name = None
    try:
        with st.spinner("Uploading video for AI analysis..."):
            video_file = genai.upload_file(video_file_path)
            uploaded_file_name = video_file.name

        processing_bar = st.progress(0, text="Processing video with AI... Please wait, this may take a moment.")
        progress_percentage = 0
        while video_file.state.name == "PROCESSING":
            time.sleep(2)
            video_file = genai.get_file(uploaded_file_name)
            progress_percentage = min(progress_percentage + 5, 99)
            processing_bar.progress(progress_percentage)

        if video_file.state.name == "FAILED":
            st.error("Failed to process video with our AI. Please try again or with a different video.")
            return "Error: Video processing failed."

        # --- Updated Prompt for Direct Content Extraction ---
        # The prompt is updated EXACTLY as requested, reinforcing no introduction.
        prompt = (
            "You are a highly precise content extraction AI. Your task is to extract the text from the video without change anything thing in the content"
            "you never should Put a introduction before the extracted text, Put Just the extracted text exactly, without change anything else"
            " the language of extracted text should be same of the video .. be carefuly on that"
            "Present the extracted content in clean, well-formatted paragraphs, ordering events chronologically as they appear in the video. "
        )
        
        # Using gemini-2.0-flash model (UNMODIFIED AS PER INSTRUCTION)
        model = genai.GenerativeModel('gemini-2.0-flash')
        response = model.generate_content([prompt, video_file], stream=False)
        
        if response and hasattr(response, 'text'):
            return response.text
        else:
            st.error("Our AI service did not return expected text content.")
            return "Error: AI analysis did not return valid text."

    except Exception as e:
        st.error(f"An error occurred during AI analysis: {e}")
        return "Error: Could not extract information using our AI service."
    finally:
        if uploaded_file_name:
            try:
                genai.delete_file(uploaded_file_name)
                st.success("Temporary file successfully cleaned up from our AI system.")
            except Exception as e:
                st.warning(f"Failed to delete uploaded file from our AI system: {e}. "
                           "Please check your usage or manually clear if possible.")

# --- Helper Function for Chat with Extracted Text ---
def chat_with_extracted_text(user_query):
    # Specific response for identity questions
    identity_keywords = ["who are you", "who developed you", "what technology", "who made you", "your developer", "your creator"]
    if any(keyword in user_query.lower() for keyword in identity_keywords):
        return "TahiriExtractor AI-chat assistance, developed by TahiriExtractor team, use LLMs technology."

    if not st.session_state.extracted_text:
        return "Please extract video content first to enable chat functionality."

    # Construct chat history for the model, focusing on the extracted text as context
    chat_model = genai.GenerativeModel('gemini-2.0-flash')
    
    # Initialize chat history with a system prompt and the extracted text as initial context
    # This ensures the model always has the extracted text as the primary context
    messages = [
        {"role": "user", "parts": [
            "You are TahiriExtractor AI-chat assistance. Your sole purpose is to answer questions strictly based on the provided video content summary. Provide professional, deep, and concise answers, focusing solely on analysis of the text. Do not answer questions outside of this context. If a user asks about your identity, development, or technology, respond with: 'TahiriExtractor AI-chat assistance, developed by TahiriExtractor team, use LLMs technology.' Do not deviate from this specific response for such queries."
            f"Here is the video content summary: {st.session_state.extracted_text}"
        ]}
    ]

    # Add the user's current query
    messages.append({"role": "user", "parts": [user_query]})

    try:
        response = chat_model.generate_content(messages)
        if response and hasattr(response, 'text'):
            return response.text
        else:
            return "Could not generate a response based on the provided text."
    except Exception as e:
        return f"An error occurred during chat: {e}"

# --- Custom CSS for Theming and Styling ---
st.markdown("""
<style>
    /* Global Inter font import and application */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    body {
        font-family: 'Inter', sans-serif;
    }

    /* Streamlit overrides for better theme compatibility */
    .stApp {
        color: var(--text-color);
        background-color: var(--background-color);
    }
    h1, h2, h3, h4, h5, h6 {
        color: var(--text-color);
    }
    p {
        color: var(--text-color);
    }

    /* --- Main Title Container --- */
    .main-title-container {
        padding: 30px 0;
        background: var(--background-color-secondary);
        border-radius: 15px;
        margin-bottom: 25px;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.1);
        text-align: center;
    }
    .main-title {
        color: #0A6EFD; /* Prominent professional blue */
        font-size: 2.8em;
        font-weight: 700;
        letter-spacing: 1px;
        margin-bottom: 10px;
        text-shadow: 1px 1px 2px rgba(0,0,0,0.05);
    }
    .main-subtitle {
        color: var(--text-color-secondary);
        font-size: 1.2em;
        max-width: 800px;
        margin: 0 auto;
        line-height: 1.6;
    }
    .info-box {
        background-color: var(--background-color-tertiary);
        padding: 20px;
        border-radius: 10px;
        text-align: center;
        margin-top: 25px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.05);
    }
    .info-box p {
        font-size: 1.05em;
        color: var(--text-color);
    }
    .info-box strong {
        color: #0A6EFD;
    }

    /* --- Main Content Wrapper (new for the "white surrounding") --- */
    .main-content-wrapper {
        background-color: var(--background-color-secondary); /* This will be light gray/dark gray */
        padding: 25px;
        border-radius: 15px;
        box-shadow: 0 6px 20px rgba(0, 0, 0, 0.15); /* More pronounced shadow */
        margin-bottom: 30px;
    }

    /* --- Extracted Text Output Styling --- */
    .extracted-text-output {
        background-color: var(--background-color-tertiary); /* Slightly different from wrapper for contrast */
        padding: 20px;
        border-radius: 12px;
        border-left: 6px solid #0A6EFD;
        margin-bottom: 25px;
        overflow-wrap: break-word;
        font-family: 'Inter', sans-serif;
        line-height: 1.6;
        color: var(--text-color);
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
    }
    
    /* --- Video Player Styling --- */
    .stVideo {
        border-radius: 12px;
        overflow: hidden;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.15);
        margin-bottom: 25px;
    }

    /* --- Chat Container for scrolling --- */
    .chat-messages-container {
        max-height: 400px;
        overflow-y: auto;
        padding: 15px;
        border-radius: 10px;
        background-color: var(--background-color-tertiary);
        box-shadow: inset 0 0 5px rgba(0,0,0,0.05);
        display: flex;
        flex-direction: column;
    }

    /* --- Chat Message Styling (ChatGPT-like) --- */
    .chat-message-user, .chat-message-ai {
        padding: 12px 18px;
        border-radius: 18px;
        margin-bottom: 8px;
        max-width: 80%;
        font-size: 0.95em;
        line-height: 1.5;
        box-shadow: 0 1px 3px rgba(0,0,0,0.08);
        word-wrap: break-word;
        word-break: break-word;
    }
    .chat-message-user {
        background-color: var(--primary-color-20);
        color: var(--text-color);
        margin-left: auto;
        border-bottom-right-radius: 4px;
    }
    .chat-message-ai {
        background-color: var(--secondary-background-color);
        color: var(--text-color);
        margin-right: auto;
        border-bottom-left-radius: 4px;
    }

    /* --- Chat Input Styling (Elastic Text Area) --- */
    .stTextArea label {
        font-weight: 600;
        color: var(--text-color);
    }
    .stTextArea [data-baseweb="textarea"] textarea {
        min-height: 50px; /* Minimum height for the input */
        max-height: 150px; /* Maximum height before scrolling */
        overflow-y: auto; /* Enable scrolling when content exceeds max-height */
        resize: vertical; /* Allow manual vertical resizing */
        border-radius: 10px;
        padding: 10px;
        background-color: var(--background-color-secondary);
        color: var(--text-color);
        font-family: 'Inter', sans-serif;
    }
    /* Hide scrollbar when not needed, but keep functionality */
    .stTextArea [data-baseweb="textarea"] textarea::-webkit-scrollbar {
        width: 8px;
    }
    .stTextArea [data-baseweb="textarea"] textarea::-webkit-scrollbar-thumb {
        background-color: rgba(0,0,0,0.2);
        border-radius: 4px;
    }
    .stTextArea [data-baseweb="textarea"] textarea::-webkit-scrollbar-track {
        background-color: transparent;
    }

    /* --- Sidebar Custom Styling --- */
    section.main[data-testid="stSidebar"] > div:first-child {
        background-color: var(--background-color-secondary); /* Grayer background, adapts to theme */
        color: var(--text-color);
        padding-top: 20px;
    }

    .sidebar .stButton > button {
        width: 100%;
        text-align: left;
        padding: 10px 15px;
        border: none;
        background-color: transparent;
        color: var(--text-color);
        font-size: 1.1em;
        font-weight: 500;
        border-radius: 8px;
        margin-bottom: 5px;
        transition: background-color 0.2s, color 0.2s;
    }
    .sidebar .stButton > button:hover {
        background-color: rgba(0, 123, 255, 0.1); /* Light blue hover */
        color: #0A6EFD; /* Vibrant blue on hover */
    }
    .sidebar .stButton > button.active-page {
        background-color: #0A6EFD;
        color: white;
        font-weight: 600;
        box-shadow: 0 2px 5px rgba(0, 123, 255, 0.3);
    }
    .sidebar h2, .sidebar h3 { /* Styling for sidebar section headers */
        color: var(--text-color); /* Ensure headers are readable on new gray bg */
        font-size: 1.25em; /* Slightly larger */
        margin-top: 25px;
        margin-bottom: 10px; /* Reduced margin */
        border-bottom: 1px solid var(--border-color); /* Subtle separator */
        padding-bottom: 8px;
        padding-left: 15px; /* Indent headers slightly */
        padding-right: 15px;
    }
    .sidebar p { /* Styling for general text in sidebar */
        font-size: 0.95em;
        line-height: 1.4;
        color: var(--text-color);
        margin-bottom: 10px;
        padding-left: 15px;
        padding-right: 15px;
    }
    .sidebar .stAlert { /* Styling for info/warning boxes in sidebar */
        font-size: 0.9em;
        padding-left: 15px;
        padding-right: 15px;
    }
</style>

<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">

<div class="main-title-container">
    <h1 class="main-title">TahiriExtractor - Video Ultra Transcription</h1>
    <p class="main-subtitle">
        Your ultimate AI-powered tool for extracting comprehensive visual information and on-screen text from videos.
    </p>
    <div class="info-box">
        <p>Harnessing <strong>advanced Artificial Intelligence</strong> to deliver detailed, readable summaries.</p>
        <p><strong>Key Point:</strong> Supports videos up to <strong>15 minutes</strong> in duration for efficient analysis. Focuses on visual content; does not transcribe audio.</p>
    </div>
</div>
""", unsafe_allow_html=True)

# --- Sidebar Content ---
with st.sidebar:
    st.markdown("<h2>Navigation</h2>", unsafe_allow_html=True)

    # Navigation buttons that control the main content area
    if st.button("Video Extraction", key="sidebar_btn_video_extraction"):
        st.session_state.main_page_selection = "Video Extraction"
        st.rerun()
    if st.button("Chat with Content", key="sidebar_btn_chat_content"):
        st.session_state.main_page_selection = "Chat with Content"
        st.rerun()
    
    # Apply active class styling to the selected button
    st.markdown(f"""
    <script>
        const currentPage = "{st.session_state.main_page_selection}";
        const videoBtn = document.querySelector('button[key="sidebar_btn_video_extraction"]');
        const chatBtn = document.querySelector('button[key="sidebar_btn_chat_content"]');
        
        if (videoBtn && currentPage === "Video Extraction") {{
            videoBtn.classList.add('active-page');
        }} else if (videoBtn) {{
            videoBtn.classList.remove('active-page');
        }}

        if (chatBtn && currentPage === "Chat with Content") {{
            chatBtn.classList.add('active-page');
        }} else if (chatBtn) {{
            chatBtn.classList.remove('active-page');
        }}
    </script>
    """, unsafe_allow_html=True)


    st.markdown("<h3>How to Use TahiriExtractor</h3>", unsafe_allow_html=True)
    st.markdown("""
    * **Upload Video:** Start by uploading your video file on the 'Video Extraction' page.
    * **Generate Summary:** Click 'Generate Content Summary' to initiate AI analysis.
    * **View & Copy:** The extracted visual text will appear directly below the video.
    * **Chat:** Switch to 'Chat with Content' to ask questions about the extracted text.
    """)

    st.markdown("<h3>About TahiriExtractor</h3>", unsafe_allow_html=True)
    st.markdown("""
    **TahiriExtractor** is an innovative application leveraging **advanced Artificial Intelligence** to:
    * Extract deep textual insights from video visual content.
    * Analyze video frames for a comprehensive, structured summary.
    * Capture and transcribe any on-screen text, actions, and visual narratives.
    * **Important Note:** This application does not transcribe audio content (speech-to-text).

    Our mission is to provide an efficient and powerful tool for researchers, content creators, and professionals
    who need to quickly understand and utilize the visual narrative embedded in video content.
    """)

    st.markdown("<h3>Contact Us</h3>", unsafe_allow_html=True)
    st.markdown("""
    Have questions, feedback, or need support? We'd love to hear from you!

    Please feel free to reach out to our team. We are committed to providing excellent support and continuously improving TahiriExtractor.

    Contact us via email:
    [TahiriExtractor.veo.net](mailto:oussama.sebrou@gmail.com?subject=Professional%20Inquiry%20regarding%20TahiriExtractor%20Application&body=Dear%20TahiriExtractor%20Team%2C%0A%0AI%20am%20writing%20to%20you%20from%20the%20TahiriExtractor%20application.%20My%20inquiry%20is%20regarding%3A%20%5Bbriefly%20state%20purpose%5D%0A%0AThank%20you%20for%20your%20time%20and%20assistance.%0A%0ASincerely%2C%0A%5BYour%20Name%5D)
    """)

# --- Main Content Area (Conditional Rendering) ---
# New wrapper for the main content to give it a distinct "white/light" background
st.markdown("<div class='main-content-wrapper'>", unsafe_allow_html=True)

if st.session_state.main_page_selection == "Video Extraction":
    st.subheader("Extract Visual Content")
    uploaded_file = st.file_uploader(
        "Upload a video file (MP4, MOV, MKV, AVI, WEBM, etc.)",
        type=["mp4", "mov", "mkv", "avi", "webm"],
        help="Our AI system supports videos up to approximately 15 minutes in length. Larger files may fail or take longer."
    )

    transcript_display_area = st.empty()

    if uploaded_file is not None:
        # Save the uploaded file to a temporary location
        with tempfile.NamedTemporaryFile(delete=False, suffix=Path(uploaded_file.name).suffix) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            video_path = tmp_file.name

        st.video(video_path)

        if st.button("Generate Content Summary", type="primary", use_container_width=True):
            with st.spinner("Analyzing video visual content with our AI... This may take a moment based on video length and complexity."):
                extracted_text = extract_text_with_ai(video_path)
                st.session_state.extracted_text = extracted_text # Store extracted text in session state
                st.session_state.chat_history = [] # Reset chat history when new text is extracted

            with transcript_display_area.container():
                st.markdown("---")
                
                # Display the extracted text directly without an introduction/subheader
                st.markdown(
                    f"""
                    <div class="extracted-text-output">
                        {extracted_text}
                    </div>
                    """,
                    unsafe_allow_html=True
                )

                # Copyable text area (with built-in copy icon)
                st.text_area(
                    "Copyable Visual Content Summary",
                    value=extracted_text,
                    height=300,
                    key="copy_summary_area",
                    help="You can easily copy the entire summary from this text box using the built-in clipboard icon (top-right of this text area)."
                )

                # Create and download Word document
                doc = Document()
                doc.add_heading('Video Visual Content Summary', level=1)
                for paragraph_text in extracted_text.split('\n\n'):
                    if paragraph_text.strip():
                        doc.add_paragraph(paragraph_text.strip())

                bio = io.BytesIO()
                doc.save(bio)
                bio.seek(0)

                st.download_button(
                    label="Download as Word (DOCX)",
                    data=bio.getvalue(),
                    file_name="video_visual_content_summary.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    help="Download the extracted visual content summary as a Microsoft Word document for offline use."
                )
        
        # Ensure temporary file is cleaned up after use
        if os.path.exists(video_path):
            os.unlink(video_path)

elif st.session_state.main_page_selection == "Chat with Content":
    st.subheader("Chat with Extracted Content")
    if not st.session_state.extracted_text:
        st.info("Please go to 'Video Extraction' to upload a video and generate content first to enable chat.")
    else:
        st.write("Ask questions about the video content that was just extracted.")

        # Chat messages container with scrolling
        st.markdown("<div id='chat-messages-scroll-container' class='chat-messages-container'>", unsafe_allow_html=True)
        for message in st.session_state.chat_history:
            if message['role'] == 'user':
                st.markdown(f"<div class='chat-message-user'><b>You:</b> {message['content']}</div>", unsafe_allow_html=True)
            else:
                st.markdown(f"<div class='chat-message-ai'><b>TahiriExtractor AI:</b> {message['content']}</div>", unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True) # Close chat messages container

        # JavaScript to scroll to the bottom of the chat container
        st.markdown("""
        <script>
            var objDiv = document.getElementById("chat-messages-scroll-container");
            if (objDiv) {
                objDiv.scrollTop = objDiv.scrollHeight;
            }
        </script>
        """, unsafe_allow_html=True)

        # Replaced st.text_input with st.text_area for elastic input
        user_chat_query = st.text_area("Your question about the content:", key="chat_input", height=68) # Added initial height

        if st.button("Ask TahiriExtractor AI", type="secondary"):
            if user_chat_query:
                st.session_state.chat_history.append({"role": "user", "content": user_chat_query})
                with st.spinner("TahiriExtractor AI is thinking..."):
                    ai_response = chat_with_extracted_text(user_chat_query)
                    st.session_state.chat_history.append({"role": "ai", "content": ai_response})
                st.rerun()
            else:
                st.warning("Please enter a question to chat.")
        
        # New feature: Copy Chat History
        if st.session_state.chat_history:
            # Format chat history for copying
            formatted_chat_history = ""
            for msg in st.session_state.chat_history:
                formatted_chat_history += f"{msg['role'].capitalize()}: {msg['content']}\n"
            
            st.text_area(
                "Copy Entire Chat History",
                value=formatted_chat_history.strip(), # .strip() removes trailing newline
                height=150,
                key="copy_chat_history_area",
                help="Copy the full conversation history from this box using the built-in clipboard icon."
            )


# Ensure the main content wrapper is closed
st.markdown("</div>", unsafe_allow_html=True)

